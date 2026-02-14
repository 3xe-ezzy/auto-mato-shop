import docx
import sys
import os

# Force UTF-8 output
sys.stdout.reconfigure(encoding='utf-8')

def read_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    print(f"--- Reading {file_path} ---")
    doc = docx.Document(file_path)
    
    print("Paragraphs:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"P{i}: {para.text}")

    print("\nTables:")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"  {row_text}")
    print("---------------------------")

if __name__ == "__main__":
    files = [
        "FAHRZEUGDATEN-Teil 1.docx",
        "xport const carData.docx",
        "from docx import Document.docx"
    ]
    for f in files:
        read_docx(f)
