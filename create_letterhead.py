from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_bg(cell, color_hex):
    """Setzt die Hintergrundfarbe einer Tabellenzelle."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Dokument erstellen
document = Document()

# Seitenränder anpassen (Standard für den Inhalt)
section = document.sections[0]
section.top_margin = Cm(5.0) # Platz für den schwarzen Balken
section.bottom_margin = Cm(5.0) # Platz für die Fußzeile
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.0)

# --- KOPFZEILE (Schwarzer Balken) ---
header = section.header
# Trick: Wir nutzen eine Tabelle für den Hintergrund
header_table = header.add_table(rows=1, cols=1, width=Cm(21))
header_table.autofit = False
header_table.allow_autofit = False
# Tabelle muss breiter als die Seite wirken für den "Randlos"-Effekt
# In Word ist das via Code schwer perfekt randlos zu kriegen ohne komplexe Hacks,
# aber das hier erstellt den schwarzen Block.
cell = header_table.cell(0, 0)
cell.height = Cm(4)
set_cell_bg(cell, "000000") # Schwarz

# Text im Header (Platzhalter für Logo)
paragraph = cell.paragraphs[0]
run = paragraph.add_run("[HIER DEIN LOGO EINFÜGEN]\n(Textfarbe Weiß wählen oder Bild über den Text legen)")
run.font.color.rgb = RGBColor(255, 255, 255)
run.font.bold = True
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- FUSSZEILE (Grauer Balken) ---
footer = section.footer
footer_table = footer.add_table(rows=1, cols=3, width=Cm(17)) # Breite angepasst an Seitenränder
footer_table.autofit = True

# Hintergrund für die ganze Zeile (simuliert durch Zellen)
for cell in footer_table.rows[0].cells:
    set_cell_bg(cell, "D9D9D9") # Hellgrau
    cell.height = Cm(4)

# Spalte 1: Adresse
c1 = footer_table.cell(0, 0)
p1 = c1.paragraphs[0]
p1.add_run("ANSCHRIFT:\n").bold = True
p1.add_run("MATO AUTOMOBILE\nAhmed Abdalla Import Export\nRödelheimer Landstraße 75\n60487 Frankfurt am Main")
p1.style.font.size = Pt(8)

# Spalte 2: Kontakt
c2 = footer_table.cell(0, 1)
p2 = c2.paragraphs[0]
p2.add_run("KONTAKT:\n").bold = True
p2.add_run("Tel: +49 69 97785893\nFax: +49 69 97785894\nMobil: +49 171 1482343\nE-Mail: info@mato-mobile.de")
p2.style.font.size = Pt(8)

# Spalte 3: Bank / Steuer
c3 = footer_table.cell(0, 2)
p3 = c3.paragraphs[0]
p3.add_run("BANKVERBINDUNG:\n").bold = True
p3.add_run("Bank: [Bankname]\nIBAN: [DE...]\nBIC: [BIC...]\n\nSteuer-Nr.: [Nummer]")
p3.style.font.size = Pt(8)

# --- INHALT ---
# Absenderzeile klein
absender = document.add_paragraph()
absender.add_run("MATO Automobile - Rödelheimer Landstraße 75 - 60487 Frankfurt am Main").font.size = Pt(7)
absender.paragraph_format.space_after = Pt(24)

# Adressfeld
document.add_paragraph("Empfänger Name\nStraße Hausnummer\nPLZ Stadt")

document.add_paragraph("\n\n")
document.add_paragraph("Datum: 29.11.2025", style='Normal').alignment = WD_ALIGN_PARAGRAPH.RIGHT

document.add_heading('Betreff: Hier Betreff eingeben', level=2)

document.add_paragraph("Sehr geehrte Damen und Herren,\n\ndies ist der Textbereich für Ihren Brief.")

# Speichern
document.save('MATO_Briefbogen.docx')
print("Datei 'MATO_Briefbogen.docx' wurde erstellt.")
